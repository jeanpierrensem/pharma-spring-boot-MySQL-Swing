#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Génère la synthèse .docx pour le bureau Documents."""
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading_custom(doc, text, level):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
    return p


def set_body_style(paragraph):
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def main():
    out = Path("/Users/nsemjean/Documents/Synthese_Techniques_Memorisation_Rapide_Armando_Elle.docx")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading(
        "Synthèse : Techniques de mémorisation rapide\nArmando Elle (volume « Memoria », tome 1)",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Document rédigé en français simple — objectif : retrouver les idées utiles, "
        "les exemples concrets et les bon pratiques sans jargon inutile."
    )

    doc.add_paragraph(
        "Note importante sur la structure : l’ouvrage d’Armando Elle est souvent publié "
        "sous le titre italien « Tecniche di Memorizzazione Veloce » et ses traductions "
        "peuvent varier. Les sommaires détaillés ne sont pas toujours disponibles "
        "publiquement. Ce document organise donc le contenu en chapitres thématiques "
        "qui correspondent aux thèmes annoncés par les éditions et aux méthodes "
        "classiques que l’auteur cite (imagination, lieux, association, nombres, listes, "
        "langues, entraînement). Si votre livre papier utilise d’autres titres de "
        "chapitres, vous pouvez simplement faire correspondre chaque partie à un thème "
        "proche."
    )

    doc.add_paragraph()

    chapters = [
        {
            "title": "Chapitre 1 — Pourquoi entraîner sa mémoire (et ce qu’on peut attendre)",
            "resume": (
                "Ce volet pose le cadre : la mémoire rapide n’est pas une « astuce magique » "
                "mais un ensemble d’habitudes et de techniques qui utilisent surtout l’image "
                "et l’ordre. Le livre insiste sur la simplicité des méthodes et sur le fait "
                "qu’elles donnent des résultats si on les pratique."
            ),
            "concepts": (
                "• Mémoire de travail vs mémoire entraînée : on retient mieux ce qui est "
                "concret, mouvementé, émotionnel et bien ordonné.\n"
                "• Principe clé : transformer l’abstrait (chiffre, mot, idée) en quelque "
                "chose de visible dans la tête.\n"
                "• La régularité compte autant que l’intelligence."
            ),
            "problemes": (
                "« Je relis dix fois et j’oublie », « Je panique en réunion », "
                "« J’apprends lentement », « Je confonds l’ordre des idées »."
            ),
            "pratiques": (
                "• Faire de courts blocs d’entraînement (souvent plus efficaces qu’une "
                "grosse séance rare).\n"
                "• Mesurer un résultat simple (ex. nombre d’éléments mémorisés, temps) "
                "pour rester motivé.\n"
                "• Commencer par des listes courtes puis augmenter la difficulté."
            ),
            "entreprise": (
                "En entreprise, une mémoire fiable réduit les erreurs (chiffres, procédures), "
                "accélère l’onboarding et rend les échanges plus clairs (on restitue un "
                "message sans hésitation)."
            ),
        },
        {
            "title": "Chapitre 2 — Attention, concentration et images « qui accrochent »",
            "resume": (
                "Avant de mémoriser vite, il faut fixer l’information. Ce chapitre développe "
                "l’idée que l’imagination est le carburant de la mémorisation : plus une image "
                "est bizarre, grande, colorée ou émotionnelle, plus elle se rappelle."
            ),
            "concepts": (
                "• L’attention sélectionne : sans focus, aucune technique ne fonctionne.\n"
                "• Exagération et mouvement : notre cerveau retient le « trop grand », "
                "le « trop petit », le « en train de tomber ».\n"
                "• Unité d’action : une scène par idée évite l’enchevêtrement."
            ),
            "problemes": (
                "Distractions, lecture passive, révisions ennuyeuses, oubli immédiat après "
                " avoir « compris »."
            ),
            "pratiques": (
                "• Rephraser l’info en une mini-scène (5–10 secondes) avant de la ranger "
                "dans un parcours mental.\n"
                "• Varier les sens : son + mouvement + couleur quand c’est possible.\n"
                "• Couper les notifications pendant un exercice."
            ),
            "entreprise": (
                "Mieux fixer un brief ou un chiffre clé évite les allers-retours par email "
                "et renforce la crédibilité lors des présentations."
            ),
        },
        {
            "title": "Chapitre 3 — La méthode des lieux (palais de la mémoire)",
            "resume": (
                "Technique historique (souvent associée à la tradition antique et aux "
                "maîtres de la mémoire) : on place chaque idée à un endroit précis d’un "
                "itinéraire connu (maison, trajet métro, bâtiment du bureau)."
            ),
            "concepts": (
                "• Lieux ordonnés : porte, couloir, table… dans un ordre fixe.\n"
                "• Ancrage : une image par lieu pour « poser » l’information.\n"
                "• Récupération : on rejoue le trajet pour retrouver la liste dans l’ordre."
            ),
            "problemes": (
                "Oublier l’ordre d’un exposé, d’une procédure, d’une checklist, d’arguments "
                "à dire en entretien."
            ),
            "pratiques": (
                "• Réutiliser 2–3 parcours maîtrisés plutôt que d’en inventer dix.\n"
                "• Ne pas surcharger un même lieu : une image forte par emplacement.\n"
                "• Répéter à voix basse le parcours le lendemain (consolidation)."
            ),
            "entreprise": (
                "Idéal pour un pitch, un audit qualité, une formation interne : vous "
                "découpez le contenu en étapes géographiques stables dans votre tête."
            ),
        },
        {
            "title": "Chapitre 4 — Mémoriser des listes, énumérations et présentations",
            "resume": (
                "Application directe du parcours mental : des points à dire, des causes "
                "« top 5 », des étapes de projet. Chaque point devient une scène mémorable "
                "liée au lieu suivant."
            ),
            "concepts": (
                "• Chaînage par ordre (lié aux lieux) plutôt qu’au hasard.\n"
                "• Clarté des symboles : une image représente un sens, pas un mot ambigu.\n"
                "• Répétition espacée : revoir la liste 10 minutes après, puis le lendemain."
            ),
            "problemes": (
                "Listes qui se mélangent, présentations improvisées qui oublient des "
                "sections, révisions interminables."
            ),
            "pratiques": (
                "• Toujours fixer le nombre d’éléments (ex. 7 points = 7 lieux).\n"
                "• Tester en fermant les notes : c’est le test qui révèle les trous.\n"
                "• Pour les listes longues : segmenter en sous-parcours (étages, zones)."
            ),
            "entreprise": (
                "Un manager qui structure sa communication en blocs ordonnés gagne du "
                "temps et rassure ses équipes (messages reproductibles)."
            ),
        },
        {
            "title": "Chapitre 5 — Nombres et chiffres : associer chaque nombre à une image",
            "resume": (
                "Le livre introduit des systèmes pour transformer les chiffres en images "
                "(souvent via des formes ou des sons) afin de mémoriser codes, dates, "
                "chiffres d’affaires, étapes numérotées."
            ),
            "concepts": (
                "• Les chiffres seuls sont abstraits ; une image « crochet » les rend "
                "mémorables.\n"
                "• On peut combiner plusieurs chiffres en une mini-histoire (souvent sur "
                "un lieu).\n"
                "• Régularité : un système personnel stable évite la confusion."
            ),
            "problemes": (
                "Confondre des codes PIN, dates, références produit, KPIs en réunion."
            ),
            "pratiques": (
                "• Choisir un petit jeu de correspondances (0–9) et s’y tenir.\n"
                "• S’entraîner d’abord sur 4 chiffres, puis 8, puis plus.\n"
                "• Réécrire de mémoire puis corriger immédiatement."
            ),
            "entreprise": (
                "Utile pour conformité, vente (prix, marges), support client (numéros de "
                "ticket) et reporting rapide sans dépendre uniquement du téléphone."
            ),
        },
        {
            "title": "Chapitre 6 — Mots, concepts et langues : du sens à l’image",
            "resume": (
                "Pour vocabulaire et notions nouvelles, l’auteur s’inspire d’approches comme "
                "l’association par mot-clé (son ressemblant) + image. L’objectif est de créer "
                "un pont entre le mot étranger et un souvenir visuel fort."
            ),
            "concepts": (
                "• Ancrage par ressemblance de son (approximatif mais mémorable).\n"
                "• Image qui raconte le sens, pas seulement la traduction littérale.\n"
                "• Petites séries + révision le jour suivant."
            ),
            "problemes": (
                "Vocabulaire qui « ne tient pas », confusion entre termes proches, ennui "
                "des listes plates."
            ),
            "pratiques": (
                "• Pour chaque mot : 10 secondes pour une image; puis insertion dans un lieu.\n"
                "• Prononcer à haute voix après l’image (liaison cerveau oreille).\n"
                "• Préférer 15 mots bien ancrés que 60 mots relus."
            ),
            "entreprise": (
                "Termes techniques, acronymes, glossaires métier : l’ancrage visuel accélère "
                "l’intégration des nouveaux arrivants et des équipes multilingues."
            ),
        },
        {
            "title": "Chapitre 7 — Séquence longue et spectacle de mémoire (ex. ordre des cartes)",
            "resume": (
                "Le manuel met en avant qu’avec entraînement on peut retenir des enchaînements "
                "longs — l’exemple typiquement cité est l’ordre d’un jeu de cartes. "
                "L’intuition pédagogique : ce n’est pas la « longueur » qui est magique, "
                "c’est la combinaison d’images + lieux + répétition."
            ),
            "concepts": (
                "• Découpage en paquets (petits groupes dans le parcours).\n"
                "• Précision : chaque carte (ou élément) a une image distinctive.\n"
                "• Vitesse croît avec l’entraînement, pas avec la pression."
            ),
            "problemes": (
                "Se dire « je n’ai pas une bonne mémoire » face à une tâche longue ; "
                "abandonner trop tôt."
            ),
            "pratiques": (
                "• Progression douce : demi-jeu avant jeu complet.\n"
                "• Repérer ses erreurs systématiques (confusion de deux images).\n"
                "• Faire une révision « à blanc » le soir même."
            ),
            "entreprise": (
                "Même principe pour enchaîner beaucoup d’informations structurées : "
                "procédures complexes, scénarios d’incident, trames de certification — "
                "la décomposition + ancrage ordonné domine la force brute."
            ),
        },
        {
            "title": "Chapitre 8 — Plan d’entraînement, erreurs fréquentes et discipline douce",
            "resume": (
                "Le livre insiste sur des exercices rapides et un style accessible : la "
                "mémoire est un muscle, mais il faut éviter la surcharge, l’improvisation "
                "sans répétition, et les images trop floues."
            ),
            "concepts": (
                "• Répétition espacée > bourrage de dernière minute.\n"
                "• Qualité d’image > quantité d’informations mal figées.\n"
                "• Auto-évaluation : se tester, pas seulement relire."
            ),
            "problemes": (
                "Démotivation, séances irrégulières, perfectionnisme (« ce n’est pas assez "
                "réaliste »)."
            ),
            "pratiques": (
                "• 10–15 minutes par jour valent mieux qu’une heure le jour avant.\n"
                "• Tenir un carnet de codes/images pour rester cohérent.\n"
                "• Célébrer les petits gains (temps, nombre d’éléments)."
            ),
            "entreprise": (
                "Crée une culture d’apprentissage continu : moins de rework, meilleure "
                "qualité de livraison, personnes plus autonomes après formation."
            ),
        },
    ]

    for ch in chapters:
        doc.add_page_break()
        add_heading_custom(doc, ch["title"], level=1)
        doc.add_paragraph().add_run("Résumé simple").bold = True
        p = _add_para(doc, ch["resume"])
        doc.add_paragraph().add_run("Concepts importants (expliqués simplement)").bold = True
        _add_para(doc, ch["concepts"])
        doc.add_paragraph().add_run("Problèmes courants que ce chapitre aide à résoudre").bold = True
        _add_para(doc, ch["problemes"])
        doc.add_paragraph().add_run("Bonnes pratiques recommandées").bold = True
        _add_para(doc, ch["pratiques"])
        doc.add_paragraph().add_run("Pourquoi c’est utile en entreprise").bold = True
        _add_para(doc, ch["entreprise"])

    doc.add_page_break()
    add_heading_custom(doc, "Fiche « astuces à retenir » (mémo express)", level=1)
    astuces = [
        "1) Abstrait → image + lieu ordonné : c’est la boîte à outils centrale du livre.",
        "2) Une idée forte par emplacement : évitez l’effet « tableau noir illisible ».",
        "3) Exagérez et animez : le mouvement et l’humour servent la mémorie, pas l’inverse.",
        "4) Testez sans support : la récupération active bat la relecture passive.",
        "5) Répétition espacée : 10 minutes après, puis le lendemain, puis selon vos échéances.",
        "6) Système stable pour les chiffres : moins on change de codes, plus on va vite.",
        "7) Progressif : listes courtes → plus long → situations réelles (réunion, exposé).",
        "8) Petit et souvent : l’entraînement régulier bat le sprint de dernière minute.",
    ]
    for a in astuces:
        doc.add_paragraph(a, style="List Number")

    doc.add_paragraph()
    doc.add_paragraph(
        "Source de référence : Armando Elle, Techniques de mémorisation rapide "
        "(traduction / série « Memoria », volume 1 — contenu souvent présenté comme manuel "
        "court avec exercices). Cette synthèse vulgarise et structure les thèmes publics ; "
        "elle ne remplace pas la lecture complète des exemples et exercices du livre."
    )

    doc.save(out)
    print(f"OK: {out}")


def _add_para(doc, text):
    p = doc.add_paragraph(text)
    set_body_style(p)
    return p


if __name__ == "__main__":
    main()
